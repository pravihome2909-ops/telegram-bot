"""
docx_generator.py
Exact TN Board format in Word.
Supports: Question Paper + Answer Key
Tamil font auto-detected.
"""

import os
import platform
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

from config import EXAM_FOOTER

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

_FONT_CANDIDATES = [
    ("/usr/share/fonts/truetype/freefont/FreeSans.ttf",           "FreeSans"),
    ("/usr/share/fonts/truetype/lohit-tamil/Lohit-Tamil.ttf",     "Lohit Tamil"),
    ("/usr/share/fonts/truetype/noto/NotoSansTamil-Regular.ttf",  "Noto Sans Tamil"),
    ("/usr/share/fonts/opentype/noto/NotoSansTamil-Regular.ttf",  "Noto Sans Tamil"),
    (os.path.join(BASE_DIR, "NotoSansTamil-Regular.ttf"),         "Noto Sans Tamil"),
    (os.path.join(BASE_DIR, "Latha.ttf"),                         "Latha"),
    (os.path.join(BASE_DIR, "FreeSans.ttf"),                      "FreeSans"),
    ("C:/Windows/Fonts/Latha.ttf",                                "Latha"),
    ("C:/Windows/Fonts/latha.ttf",                                "Latha"),
]

def _detect_font():
    for path, name in _FONT_CANDIDATES:
        if os.path.exists(path) and os.path.getsize(path) > 5000:
            print(f"[DOCX] Tamil font: {name}")
            return name
    print("[DOCX] No Tamil font — using Arial")
    return "Arial"

TAMIL_FONT = _detect_font()
ROMAN = {1: "I", 2: "II", 3: "III", 4: "IV"}


# ── Language detection ────────────────────────────────────────────────

def _is_tamil(text: str) -> bool:
    return any('\u0B80' <= ch <= '\u0BFF' for ch in text)

def _detect_language(questions: list) -> str:
    for q in questions:
        t = q.get("question", "")
        if t.strip():
            return "ta" if _is_tamil(t) else "en"
    return "en"

_HEADINGS = {
    "ta": {
        "sec1_main": "அனைத்து வினாக்களுக்கும் விடையளிக்கவும்.",
        "sec1_sub":  "",
        "sec2_main": "எவையேனும் 7 வினாக்களுக்கு விடையளிக்கவும்.",
        "sec2_sub":  "வினா எண். {q} க்கு கட்டாயம் விடையளிக்கவும்.",
        "sec3_main": "எவையேனும் 7 வினாக்களுக்கு விடையளிக்கவும்.",
        "sec3_sub":  "வினா எண். {q} க்கு கட்டாயம் விடையளிக்கவும்.",
        "sec4_main": "அனைத்து வினாக்களுக்கும் விடையளிக்கவும்.",
        "sec4_sub":  "",
        "or":        "(அல்லது)",
        "duration":  "நேரம் : 3.00 மணி",
        "total":     "மதிப்பெண்கள் : {t}",
        "answer_key":"விடை குறிப்பு",
        "answer_lbl":"விடை",
    },
    "en": {
        "sec1_main": "Answer all questions.",
        "sec1_sub":  "",
        "sec2_main": "Answer any 7 questions.",
        "sec2_sub":  "Question No. {q} is compulsory.",
        "sec3_main": "Answer any 7 questions.",
        "sec3_sub":  "Question No. {q} is compulsory.",
        "sec4_main": "Answer all questions.",
        "sec4_sub":  "",
        "or":        "(OR)",
        "duration":  "Time : 3.00 Hours",
        "total":     "Total Marks : {t}",
        "answer_key":"ANSWER KEY",
        "answer_lbl":"Ans",
    },
}

def _score(mark, count):
    if mark == 1: return count
    if mark == 2: return min(count,7)*2
    if mark == 3: return min(count,7)*3
    return count * mark


# ── XML helpers ───────────────────────────────────────────────────────

def _f(run, name, size, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(
            int(color[0:2],16), int(color[2:4],16), int(color[4:6],16))
    rPr = run._r.get_or_add_rPr()
    rf  = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rPr.insert(0, rf)
    for attr in ["w:ascii","w:hAnsi","w:cs","w:eastAsia"]:
        rf.set(qn(attr), name)

def _sp(para, before=0, after=3, ls=1.15):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    fmt.line_spacing = Pt(11 * ls)

def _hline(para, color="000000", thick=4):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(thick * 8))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

def _no_border(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBdr")
    for s in ["top","left","bottom","right","insideH","insideV"]:
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"),"none"); e.set(qn("w:sz"),"0")
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),"auto")
        bdr.append(e)
    tcPr.append(bdr)

def _no_tbl_border(tbl):
    tPr = tbl._tbl.find(qn("w:tblPr"))
    if tPr is None:
        tPr = OxmlElement("w:tblPr"); tbl._tbl.insert(0, tPr)
    tb = OxmlElement("w:tblBorders")
    for s in ["top","left","bottom","right","insideH","insideV"]:
        e = OxmlElement(f"w:{s}")
        e.set(qn("w:val"),"none"); e.set(qn("w:sz"),"0")
        e.set(qn("w:color"),"auto"); tb.append(e)
    tPr.append(tb)

def _footer(section, fn, wm_text="EduPulse-JB"):
    fp = (section.footer.paragraphs[0]
          if section.footer.paragraphs
          else section.footer.add_paragraph())
    fp.clear(); _sp(fp, 0, 0)
    r1 = fp.add_run(EXAM_FOOTER); _f(r1, fn, 9, color="333333")
    fp.add_run("\t")
    r2 = fp.add_run(wm_text); _f(r2, "Arial", 7, color="BBBBBB")
    pPr  = fp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),"right"); tab.set(qn("w:pos"),"9072")
    tab.set(qn("w:leader"),"none")
    tabs.append(tab); pPr.append(tabs)

def _sec_row(doc, fn, roman, main, sub, marks_str):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    cells = tbl.rows[0].cells
    _no_border(cells[0]); _no_border(cells[1])
    p0 = cells[0].paragraphs[0]; p0.clear(); _sp(p0, before=8, after=1)
    r = p0.add_run(f"{roman}   {main}"); _f(r, fn, 10, bold=True)
    p1 = cells[1].paragraphs[0]; p1.clear()
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT; _sp(p1, before=8, after=1)
    r2 = p1.add_run(marks_str); _f(r2, fn, 10, bold=True)
    _no_tbl_border(tbl)
    if sub:
        p2 = doc.add_paragraph(); _sp(p2, before=0, after=2)
        p2.paragraph_format.left_indent = Cm(0.5)
        r3 = p2.add_run(sub); _f(r3, fn, 9, color="333333")


# ── Core document builder ─────────────────────────────────────────────

def _build_doc(class_, subject, lesson, groups, part_d_qs,
               pairs, total, school_name, test_name,
               fn, H, is_answer_key=False,
               start_number=1, wm_text="EduPulse-JB") -> Document:
    doc = Document()
    LS  = 1.15

    for sec in doc.sections:
        sec.page_width    = Cm(21.0)
        sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(1.9)
        sec.bottom_margin = Cm(1.9)
        _footer(sec, fn, wm_text)

    # ── Header ───────────────────────────────────────────────────────
    if school_name or test_name:
        top = school_name if school_name else test_name
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 2, LS); r = p.add_run(top); _f(r, fn, 16, bold=True)
    if school_name and test_name:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 3, LS); r = p.add_run(test_name); _f(r, fn, 12)

    if is_answer_key:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _sp(p, 0, 3, LS)
        r = p.add_run(f"[ {H['answer_key']} ]")
        _f(r, fn, 12, bold=True, color="C0392B")

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(p, 0, 2, LS)
    r = p.add_run(f"{class_} – ஆம் வகுப்பு          {subject}")
    _f(r, fn, 11)

    ht = doc.add_table(rows=1, cols=3)
    ht.style = "Table Grid"
    hcells = ht.rows[0].cells
    hdata  = [
        (H["duration"],               WD_ALIGN_PARAGRAPH.LEFT),
        (lesson,                       WD_ALIGN_PARAGRAPH.CENTER),
        (H["total"].format(t=total),   WD_ALIGN_PARAGRAPH.RIGHT),
    ]
    for cell, (txt, aln) in zip(hcells, hdata):
        _no_border(cell)
        cp = cell.paragraphs[0]; cp.clear()
        cp.alignment = aln; _sp(cp, 1, 3, LS)
        r = cp.add_run(txt); _f(r, fn, 9)
    _no_tbl_border(ht)

    hr = doc.add_paragraph(); _sp(hr, 2, 4, LS)
    _hline(hr, color="000000", thick=1)

    # ── Sections ─────────────────────────────────────────────────────
    global_n = start_number

    def _q(text, indent=0.5):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        _sp(p, 1, 3, LS)
        r = p.add_run(text); _f(r, fn, 10)

    def _ans(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        _sp(p, 0, 4, LS)
        r = p.add_run(f"{H['answer_lbl']}: {text}")
        _f(r, fn, 9, color="1A5276")

    # Section I
    if 1 in groups and groups[1]:
        qs = groups[1]; n1 = len(qs)
        _sec_row(doc, fn, ROMAN[1], H["sec1_main"], H["sec1_sub"],
                 f"{n1}X1={n1}")
        for q in qs:
            _q(f"{global_n}.  {q.get('question','')}")
            opts = q.get("options","")
            if opts:
                opt_list = [o.strip() for o in opts.split("|")]
                labels   = ["(அ)","(ஆ)","(இ)","(ஈ)"]
                line = "   ".join(
                    f"{labels[i]} {opt_list[i]}"
                    for i in range(min(len(labels),len(opt_list))))
                _q(line, indent=1.0)
            if is_answer_key:
                ans = q.get("answer","") or q.get("correct_option","")
                if ans: _ans(ans)
            global_n += 1
        doc.add_paragraph()

    # Section II
    if 2 in groups and groups[2]:
        qs = groups[2]; n2 = len(qs); ans2 = min(n2,7)
        cq = global_n + ans2 - 1
        _sec_row(doc, fn, ROMAN[2], H["sec2_main"],
                 H["sec2_sub"].format(q=cq), f"{ans2}X2={ans2*2}")
        for q in qs:
            _q(f"{global_n}.  {q.get('question','')}")
            if is_answer_key:
                ans = q.get("answer","")
                if ans: _ans(ans)
            global_n += 1
        doc.add_paragraph()

    # Section III
    if 3 in groups and groups[3]:
        qs = groups[3]; n3 = len(qs); ans3 = min(n3,7)
        cq = global_n + ans3 - 1
        _sec_row(doc, fn, ROMAN[3], H["sec3_main"],
                 H["sec3_sub"].format(q=cq), f"{ans3}X3={ans3*3}")
        for q in qs:
            _q(f"{global_n}.  {q.get('question','')}")
            if is_answer_key:
                ans = q.get("answer","")
                if ans: _ans(ans)
            global_n += 1
        doc.add_paragraph()

    # Section IV
    if part_d_qs and pairs > 0:
        _sec_row(doc, fn, ROMAN[4], H["sec4_main"], H["sec4_sub"],
                 f"{pairs}X5={pairs*5}")
        for i in range(0, len(part_d_qs)-1, 2):
            qa  = part_d_qs[i].get("question","")
            qb  = part_d_qs[i+1].get("question","")
            ans = part_d_qs[i].get("answer","")
            _q(f"{global_n}.  {qa}")
            por = doc.add_paragraph()
            por.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _sp(por, 1, 1, LS)
            r = por.add_run(H["or"]); _f(r, fn, 10)
            _q(f"          {qb}")
            if is_answer_key and ans:
                _ans(ans)
            doc.add_paragraph()
            global_n += 1

    return doc, global_n


# ── Public: Question Paper DOCX ──────────────────────────────────────

def generate_question_paper_docx(
        class_: str, subject: str, lesson: str,
        questions: list, teacher_name: str = "",
        output_path: str = None,
        part_d: list = None,
        test_name: str = "",
        school_name: str = "",
        start_number: int = 1) -> str:

    if not DOCX_OK:
        print("[DOCX] python-docx not installed."); return ""
    if output_path is None:
        os.makedirs("generated_papers", exist_ok=True)
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = subject.replace(" ", "_")
        output_path = os.path.join(
            "generated_papers", f"Class{class_}_{safe}_{ts}.docx")

    fn        = TAMIL_FONT
    all_qs    = questions + (part_d or [])
    lang      = _detect_language(all_qs)
    H         = _HEADINGS[lang]
    part_d_qs = part_d if part_d else []
    other_qs  = [q for q in questions if q.get("marks") != 5]
    groups    = {}
    for q in other_qs:
        groups.setdefault(q.get("marks",1),[]).append(q)
    pairs = len(part_d_qs) // 2
    total = (_score(1,len(groups.get(1,[]))) +
             _score(2,len(groups.get(2,[]))) +
             _score(3,len(groups.get(3,[]))) +
             pairs * 5)
    try:
        doc, _ = _build_doc(
            class_, subject, lesson, groups, part_d_qs,
            pairs, total, school_name, test_name, fn, H,
            is_answer_key=False, start_number=start_number)
        doc.save(output_path)
        print(f"[DOCX] Question paper: {output_path}")
        return output_path
    except Exception as e:
        print(f"[DOCX] Failed: {e}")
        import traceback; traceback.print_exc()
        return ""


# ── Public: Answer Key DOCX ──────────────────────────────────────────

def generate_answer_key_docx(
        class_: str, subject: str, lesson: str,
        questions: list, teacher_name: str = "",
        output_path: str = None,
        part_d: list = None,
        test_name: str = "",
        school_name: str = "",
        start_number: int = 1) -> str:

    if not DOCX_OK:
        print("[DOCX] python-docx not installed."); return ""
    if output_path is None:
        os.makedirs("generated_papers", exist_ok=True)
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = subject.replace(" ", "_")
        output_path = os.path.join(
            "generated_papers",
            f"Class{class_}_{safe}_{ts}_ANSWER_KEY.docx")

    fn        = TAMIL_FONT
    all_qs    = questions + (part_d or [])
    lang      = _detect_language(all_qs)
    H         = _HEADINGS[lang]
    part_d_qs = part_d if part_d else []
    other_qs  = [q for q in questions if q.get("marks") != 5]
    groups    = {}
    for q in other_qs:
        groups.setdefault(q.get("marks",1),[]).append(q)
    pairs = len(part_d_qs) // 2
    total = (_score(1,len(groups.get(1,[]))) +
             _score(2,len(groups.get(2,[]))) +
             _score(3,len(groups.get(3,[]))) +
             pairs * 5)
    try:
        doc, _ = _build_doc(
            class_, subject, lesson, groups, part_d_qs,
            pairs, total, school_name, test_name, fn, H,
            is_answer_key=True, start_number=start_number,
            wm_text="Answer Key — EduPulse-JB")
        doc.save(output_path)
        print(f"[DOCX] Answer key: {output_path}")
        return output_path
    except Exception as e:
        print(f"[DOCX] Answer key failed: {e}")
        import traceback; traceback.print_exc()
        return ""
